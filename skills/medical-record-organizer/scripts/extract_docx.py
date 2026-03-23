#!/usr/bin/env python3
"""
extract_docx.py — Extract text from .docx Word files.

Usage:
    python3 extract_docx.py PATH

Output JSON:
    {"success": true, "text": "...", "paragraphs": 42}
    {"success": false, "error": "..."}
"""

import json
import sys
from pathlib import Path


def extract_docx(file_path: Path) -> dict:
    try:
        from docx import Document
    except ImportError:
        return {"success": False, "error": "python-docx not installed. Run: uv pip install python-docx"}

    doc = Document(str(file_path))
    parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract table cells
    for table in doc.tables:
        for row in table.rows:
            row_parts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_parts.append(cell_text)
            if row_parts:
                parts.append("\t".join(row_parts))

    full_text = "\n".join(parts)
    return {
        "success": True,
        "text": full_text,
        "paragraphs": len(doc.paragraphs),
    }


def main():
    if len(sys.argv) < 2:
        print(json.dumps({"success": False, "error": "Usage: python3 extract_docx.py PATH"}))
        sys.exit(1)

    file_path = Path(sys.argv[1]).expanduser().resolve()

    if not file_path.exists():
        print(json.dumps({"success": False, "error": f"File not found: {file_path}"}))
        sys.exit(1)

    suffix = file_path.suffix.lower()
    if suffix == ".doc":
        print(json.dumps({
            "success": False,
            "error": "旧版 .doc 格式不支持直接提取。请在 Word 中另存为 .docx 后重试。",
        }))
        sys.exit(1)

    if suffix != ".docx":
        print(json.dumps({"success": False, "error": f"不支持的文件格式: {suffix}，仅支持 .docx"}))
        sys.exit(1)

    result = extract_docx(file_path)
    print(json.dumps(result, ensure_ascii=False))
    if not result["success"]:
        sys.exit(1)


if __name__ == "__main__":
    main()
